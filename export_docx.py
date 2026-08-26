"""Build a Word (.docx) KEMBARAN I document from selected export data."""
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips


BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_FILL = "FFD966"  # Gold, Accent 4, Lighter 40%
FIXED_TITLE = "SENARAI LEMBARAN DAN JENIS PETA"
DEFAULT_SUBTITLE = "EKSESAIS"

# Column widths from official KEMBARAN I - GIS INFO.docx (tblGrid twips)
TOPO_COL_TWIPS = (882, 1778, 3616, 1559, 1545)
LAND_DTED_COL_TWIPS = (850, 6517, 1700)


def _set_run_font(run, *, bold=False, italic=False, underline=False, size=12):
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    # Ensure East Asian / complex script also use TNR when Word opens the file
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def _add_centered_line(doc, text, *, bold=True, underline=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    _set_run_font(run, bold=bold, underline=underline)
    return para


def _add_section_title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    _set_run_font(run, bold=True, italic=True, underline=True)


def _set_cell_shading(cell, fill=HEADER_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text, *, bold=False, center=True, header=False):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run("" if text is None else str(text))
    _set_run_font(run, bold=bold)
    if header:
        _set_cell_shading(cell)


def _set_table_column_widths(table, widths_twips):
    """Apply preferred column widths to match the official kembaran layout."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(widths_twips):
                break
            cell.width = Twips(widths_twips[idx])


def _fill_table(table, headers, rows, total_label, left_align_cols=()):
    left_align_cols = set(left_align_cols)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        _set_cell_text(hdr[i], header, bold=True, header=True, center=i not in left_align_cols)

    for r_idx, row_vals in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row_vals):
            _set_cell_text(cells[c_idx], value, bold=False, center=c_idx not in left_align_cols)

    total_row = table.rows[-1].cells
    for c_idx in range(len(headers) - 1):
        _set_cell_text(total_row[c_idx], "TOTAL", bold=True)
    _set_cell_text(total_row[-1], total_label, bold=True)


def _add_data_table(doc, headers, rows, total_label, col_widths_twips, left_align_cols=()):
    table = doc.add_table(rows=len(rows) + 2, cols=len(headers))
    table.style = "Table Grid"
    _fill_table(table, headers, rows, total_label, left_align_cols=left_align_cols)
    _set_table_column_widths(table, col_widths_twips)
    return table


def build_export_docx(report_title, report_ref, selections):
    """Return an in-memory .docx matching KEMBARAN I layout."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.text = ""
    run = header_para.add_run("KEMBARAN I")
    _set_run_font(run, bold=False, underline=False)

    _add_centered_line(doc, FIXED_TITLE, bold=True, underline=True)
    _add_centered_line(
        doc,
        (report_title or DEFAULT_SUBTITLE).strip() or DEFAULT_SUBTITLE,
        bold=True,
        underline=True,
    )

    topo = selections.get("topography") or []
    sjungu = selections.get("sjungu") or []
    section_idx = 0

    def next_label(name):
        nonlocal section_idx
        letter = chr(ord("A") + section_idx)
        section_idx += 1
        return f"{letter}. {name}"

    if topo or sjungu:
        _add_section_title(doc, next_label("Raster Topography"))
        rows = []
        num = 0
        for r in topo:
            num += 1
            rows.append([
                num,
                r.get("sheetNum", ""),
                r.get("sheetName", ""),
                r.get("sheetScale", ""),
                r.get("release_year", ""),
            ])
        for r in sjungu:
            num += 1
            rows.append([
                num,
                r.get("sheetNum", ""),
                r.get("sheetName", ""),
                r.get("sheetScale", ""),
                "",
            ])
        if topo and sjungu:
            total_label = f"{len(topo)} + {len(sjungu)}"
        else:
            total_label = str(len(topo) or len(sjungu))
        _add_data_table(
            doc,
            ["NUM.", "SHEET NUM.", "SHEET NAME", "SHEET SCALE", "RELEASE YEAR"],
            rows,
            total_label,
            TOPO_COL_TWIPS,
        )

    land = selections.get("landused") or []
    if land:
        _add_section_title(doc, next_label("Landused"))
        rows = [
            [idx + 1, r.get("category", ""), idx + 1]
            for idx, r in enumerate(land)
        ]
        _add_data_table(
            doc,
            ["NUM.", "CATEGORY", "LANDUSED ID"],
            rows,
            str(len(land)),
            LAND_DTED_COL_TWIPS,
            left_align_cols=(1,),
        )

    dted = selections.get("dted") or []
    if dted:
        _add_section_title(doc, next_label("Digital Terrain Elevation Data (DTED)"))
        rows = [
            [idx + 1, r.get("id_name", ""), r.get("level", "")]
            for idx, r in enumerate(dted)
        ]
        _add_data_table(
            doc,
            ["NUM.", "IDENTIFICATION NAME", "LEVEL"],
            rows,
            str(len(dted)),
            LAND_DTED_COL_TWIPS,
            left_align_cols=(1,),
        )

    # report_ref kept out of the Word body to match the official kembaran
    _ = report_ref

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
